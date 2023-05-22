import docx

doc = docx.Document('certificate_form.docx')

style = doc.styles['Normal']
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph()
run = para.add_run('\n')
run = para.add_run('              No 2020-0001\n')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n')
run = para.add_run('Certificate')
run.bold = True
run.font.size = docx.shared.Pt(40)
para.alignment = 1 # for left, 1 for center, 2 right, 3 justify ....

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('        N a m e : Kate\n')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        Birth : 2017.12.12\n')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        Course : Python Basic\n')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        D a t e : 2021.08.05~2021.09.09\n')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('        This person has certificate\n')
run.font.size = docx.shared.Pt(20)
run = para.add_run('        graduation of Python basic course.\n')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('2021.09.19')
run.font.size = docx.shared.Pt(20)
para.alignment = 1

para = doc.add_paragraph()
run = para.add_run('\n')
run = para.add_run('Principle of \nPython academy')
run.bold = True
run.font.size = docx.shared.Pt(20)
para.alignment = 1

doc.save('certificate.docx')
