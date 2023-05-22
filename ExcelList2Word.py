from openpyxl import load_workbook
import docx

def certificate_word(name, birth, code) :
    doc = docx.Document('certificate_form.docx')

    style = doc.styles['Normal']
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run(f'              No {code}\n')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run('Certificate')
    run.bold = True
    run.font.size = docx.shared.Pt(40)
    para.alignment = 1  # for left, 1 for center, 2 right, 3 justify ....

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(f'        N a m e : {name}\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(f'        Birth : {birth}\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        Course : Python Basic\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        D a t e : 2021.08.05~2021.09.09\n')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('        This person has certificate\n')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        graduation of Python basic course.\n')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run('2021.09.19')
    run.font.size = docx.shared.Pt(20)
    para.alignment = 1

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run('Principle of \nPython academy')
    run.bold = True
    run.font.size = docx.shared.Pt(20)
    para.alignment = 1

    doc.save(f'certificate_{name}.docx')

load_wb = load_workbook('students_list.xlsx')
load_ws = load_wb.active

#name_list = []
#birthday_list = []
#code_list = []
for i in range(1, load_ws.max_row + 1) :
    #name_list.append(load_ws.cell(i, 1).value)
    #birthday_list.append(load_ws.cell(i, 2).value)
    #code_list.append(load_ws.cell(i, 3).value)
    name = load_ws.cell(i, 1).value
    birth = load_ws.cell(i, 2).value
    code = load_ws.cell(i, 3).value
    certificate_word(name, birth, code)

#print(name_list)
#print(birthday_list)
#print(code_list)